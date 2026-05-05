"""
Resume builder - .docx renderer.

Sizing & formatting (per final spec):
  NAME:          bold, 16pt, centered
  SECTION HEADS: bold, 14pt  (Summary, Technical Skills, Professional Experience,
                               Academic Projects, Education, Certifications)
  SUBHEADINGS:   bold, 12pt  (role title, project title, skill category labels)
  BODY TEXT:     11pt, not bold
  CONTACT:       10pt, not bold

Formatting details:
  - Section headings: Title Case, bold, 14pt, thin grey rule below
  - Subheadings: bold, 12pt, space before each new block
  - Skills: each line bold 12pt label + 11pt items, ends with full stop
  - Projects: colon appended to title, 4pt spacer between blocks
  - Education: institution = bold 12pt; degree lines = 11pt
  - Certifications: bulleted, 11pt
  - Name: bold
  - Contact: LinkedIn / GitHub / Portfolio as black hyperlinks (no underline)
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

import config
from models import TailoredResume

log = logging.getLogger("resume_builder")

# ── Design tokens ──────────────────────────────────────────────────────────────
FONT_NAME        = "Calibri"
NAME_SIZE        = 16       # bold
SECTION_SIZE     = 14       # bold - all section headings equal
SUBHEADING_SIZE  = 12       # bold - role titles, project titles, skill labels
BODY_SIZE        = 11       # not bold - all body text equal
CONTACT_SIZE     = 10       # not bold
MARGIN_IN        = 0.55
BLACK            = RGBColor(26, 26, 26)
BLACK_HEX        = "1A1A1A"
RULE_COLOR       = "BBBBBB"
BULLET           = "\u2022"

CONTACT_LINKS = {
    "LinkedIn":  config.LINKEDIN_URL,
    "GitHub":    config.GITHUB_URL,
    "Portfolio": config.PORTFOLIO_URL,
}


# ── Entry point ────────────────────────────────────────────────────────────────

def build_docx(resume: TailoredResume, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_margins(doc)
    _clear_styles(doc)

    _name_line(doc, resume.name)
    _contact_line(doc, resume.contact)

    _section(doc, "Summary",                 resume.summary,
             lambda d, c: _render_paragraph(d, c))
    _section(doc, "Technical Skills",        resume.skills,
             lambda d, c: _render_skills(d, c))
    _section(doc, "Professional Experience", resume.experience,
             lambda d, c: _render_narrative(d, c, colon=False))
    _section(doc, "Academic Projects",       resume.projects,
             lambda d, c: _render_narrative(d, c, colon=True))
    _section(doc, "Education",               resume.education,
             lambda d, c: _render_education(d, c))

    if resume.certifications and resume.certifications.strip():
        _section(doc, "Certifications", resume.certifications,
                 lambda d, c: _render_certifications(d, c))

    doc.save(str(output_path))
    log.info(f"Saved -> {output_path}")
    return output_path


# ── Document setup ─────────────────────────────────────────────────────────────

def _set_margins(doc: Document) -> None:
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(MARGIN_IN)


def _clear_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name  = FONT_NAME
    style.font.size  = Pt(BODY_SIZE)
    style.font.bold  = False
    style.font.color.rgb = BLACK
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)


# ── Spacing helpers ────────────────────────────────────────────────────────────

def _body_sp(p, before: float = 0, after: float = 2) -> None:
    pf = p.paragraph_format
    pf.space_before      = Pt(before)
    pf.space_after       = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing      = Pt(BODY_SIZE * 1.15)


def _head_sp(p, size: float, before: float = 0, after: float = 2) -> None:
    pf = p.paragraph_format
    pf.space_before      = Pt(before)
    pf.space_after       = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    pf.line_spacing      = Pt(size)


def _spacer(doc: Document, pts: float = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)


# ── Run helpers ────────────────────────────────────────────────────────────────

def _run(p, text: str, size: float = BODY_SIZE, bold: bool = False) -> None:
    run = p.add_run(text)
    run.font.name      = FONT_NAME
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.color.rgb = BLACK


def _hyperlink_run(paragraph, text: str, url: str, size: float = CONTACT_SIZE) -> None:
    """Black hyperlink - no underline, no italic, no bold."""
    if not url:
        _run(paragraph, text, size)
        return
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), FONT_NAME)
        rFonts.set(qn("w:hAnsi"), FONT_NAME)
        rPr.append(rFonts)

        for tag in ("w:sz", "w:szCs"):
            el = OxmlElement(tag)
            el.set(qn("w:val"), str(int(size * 2)))
            rPr.append(el)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), BLACK_HEX)
        rPr.append(color)

        u = OxmlElement("w:u")
        u.set(qn("w:val"), "none")
        rPr.append(u)

        i = OxmlElement("w:i")
        i.set(qn("w:val"), "0")
        rPr.append(i)

        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        if text.startswith(" ") or text.endswith(" "):
            t.set(qn("xml:space"), "preserve")
        r.append(t)

        hyperlink.append(r)
        paragraph._p.append(hyperlink)
    except Exception as _e:
        log.debug(f"Hyperlink creation failed for '{text}': {_e} - using plain text.")
        _run(paragraph, text, size)


# ── Header ─────────────────────────────────────────────────────────────────────

def _name_line(doc: Document, name: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _head_sp(p, NAME_SIZE, before=0, after=2)
    _run(p, name.upper(), NAME_SIZE, bold=True)


def _contact_line(doc: Document, contact: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _body_sp(p, before=0, after=5)

    parts = [pt.strip() for pt in contact.split("|")]
    for i, part in enumerate(parts):
        if i > 0:
            _run(p, " | ", CONTACT_SIZE)
        url = CONTACT_LINKS.get(part, "")
        if url or part in CONTACT_LINKS:
            _hyperlink_run(p, part, url, CONTACT_SIZE)
        else:
            _run(p, part, CONTACT_SIZE)


# ── Section frame ──────────────────────────────────────────────────────────────

def _section(doc: Document, heading: str, content: str, renderer) -> None:
    if not content or not content.strip():
        return
    _section_heading(doc, heading)
    renderer(doc, content)


def _section_heading(doc: Document, text: str) -> None:
    """Bold 14pt Title Case heading with thin grey rule. All sections same size."""
    p = doc.add_paragraph()
    _head_sp(p, SECTION_SIZE, before=6, after=2)
    _run(p, _title_case(text), SECTION_SIZE, bold=True)

    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), RULE_COLOR)
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Renderers ──────────────────────────────────────────────────────────────────

def _render_paragraph(doc: Document, content: str) -> None:
    """Plain body text (Summary). 11pt, not bold."""
    for block in _split_blocks(content):
        if not block.strip():
            continue
        p = doc.add_paragraph()
        _body_sp(p)
        _run(p, _clean_md(block.strip()), BODY_SIZE)


def _render_skills(doc: Document, content: str) -> None:
    """
    Skills block. Each line:
      [Bold 12pt label]: [11pt items].
    Ends with a full stop.
    """
    lines = [l.strip() for l in content.splitlines() if l.strip()]
    for line in lines:
        line = _clean_md(line)
        # Ensure full stop at end
        if line and not line.endswith("."):
            line = line + "."

        p = doc.add_paragraph()
        _body_sp(p, before=0, after=2)

        if ":" in line:
            label, _, rest = line.partition(":")
            # Bold label
            _run(p, label.strip() + ":", SUBHEADING_SIZE, bold=True)
            # Normal items
            _run(p, " " + rest.strip(), BODY_SIZE)
        else:
            _run(p, line, BODY_SIZE)


def _render_narrative(doc: Document, content: str, colon: bool = False) -> None:
    """
    Experience and projects.
    Processes content LINE BY LINE (not block-by-block) so heading detection
    always works correctly regardless of whether the LLM uses blank lines
    between projects or not.

    - Heading lines (project/role titles): bold 12pt, colon appended for projects
    - Body lines: 11pt, not bold
    - 6pt spacer added BETWEEN project body and next project heading only
      (never after the last project - that creates unwanted gaps before Education)
    """
    # Always split into individual lines; ignore blank lines.
    # Heading vs body is detected by line content, not position.
    lines = [_clean_md(l.strip()) for l in content.splitlines() if l.strip()]

    for idx, line in enumerate(lines):
        if _is_heading(line):
            text = line
            if colon and " | " not in line and not line.endswith(":"):
                text = line + ":"
            p = doc.add_paragraph()
            _head_sp(p, SUBHEADING_SIZE, before=0, after=1)
            _run(p, text, SUBHEADING_SIZE, bold=True)
        else:
            # Detect bullet lines (start with "- ") in both experience and projects
            if line.startswith("- "):
                text = BULLET + "  " + line[2:].strip()
                p = doc.add_paragraph()
                _body_sp(p, before=0, after=2)
                _run(p, text, BODY_SIZE)
            else:
                p = doc.add_paragraph()
                _body_sp(p, before=0, after=2)
                _run(p, line, BODY_SIZE)

            # No spacer between projects - they flow directly one after another.


def _render_education(doc: Document, content: str) -> None:
    """
    Education section.
    Institution name: bold 12pt.
    Degree / honors: 11pt not bold.
    4pt spacer between institutions.
    Robust grouping: short city/country lines merged into institution name.
    """
    raw_blocks = [b.strip() for b in re.split(r"\n\n+", content.strip()) if b.strip()]

    if len(raw_blocks) == 1:
        lines = [l.strip() for l in raw_blocks[0].splitlines() if l.strip()]
        raw_blocks = _group_education_lines(lines)

    for i, block in enumerate(raw_blocks):
        block_lines = [l.strip() for l in block.splitlines() if l.strip()]
        for j, line in enumerate(block_lines):
            line = _clean_md(line)
            if j == 0:
                p = doc.add_paragraph()
                _head_sp(p, SUBHEADING_SIZE, before=0, after=1)
                _run(p, line, SUBHEADING_SIZE, bold=True)
            else:
                p = doc.add_paragraph()
                _body_sp(p, before=0, after=2)
                _run(p, line, BODY_SIZE)


def _render_certifications(doc: Document, content: str) -> None:
    """Bulleted certifications list. 11pt."""
    lines = [l.strip() for l in content.splitlines() if l.strip()]
    for line in lines:
        line = re.sub(r"^[•\-\*▸–\d\.]+\s*", "", line).strip()
        if not line:
            continue
        if not line.endswith("."):
            line = line + "."
        p = doc.add_paragraph()
        _body_sp(p, before=0, after=2)
        _run(p, f"{BULLET}  {line}", BODY_SIZE)


# ── Helpers ────────────────────────────────────────────────────────────────────

def _title_case(text: str) -> str:
    return " ".join(w[0].upper() + w[1:] if w else w for w in text.split())


def _clean_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__",      r"\1", text)
    return text


def _split_blocks(content: str) -> list[str]:
    normalised = re.sub(r"\n{2,}", "\n\n", content.strip())
    blocks = [b.strip() for b in normalised.split("\n\n") if b.strip()]
    if len(blocks) == 1:
        blocks = [l.strip() for l in blocks[0].splitlines() if l.strip()]
    return blocks


def _is_heading(line: str) -> bool:
    if " | " in line:
        return True
    if re.search(r"\b(19|20)\d{2}$", line) and len(line) <= 90:
        return True
    if len(line) <= 70 and not re.search(r"[.!?]$", line) and not line.startswith("I "):
        return True
    return False


def _looks_like_institution(line: str) -> bool:
    if re.search(r"[.!?]$", line):  return False
    if line.startswith("I "):        return False
    if " | " in line:                return False
    degree_words = ("Bachelor", "Master", "Doctor", "PhD", "BSc", "MSc",
                    "B.S", "M.S", "GPA", "Recognition")
    if any(line.startswith(d) for d in degree_words): return False
    if len(line) <= 25:              return False
    return len(line) <= 80


def _group_education_lines(lines: list[str]) -> list[str]:
    degree_starts = ("Bachelor", "Master", "Doctor", "PhD", "BSc", "MSc",
                     "B.S", "M.S", "GPA", "Recognition")
    blocks: list[list[str]] = []
    current: list[str] = []

    for line in lines:
        is_degree     = any(line.startswith(d) for d in degree_starts)
        is_very_short = len(line) <= 25
        has_pipe      = " | " in line
        has_sent_end  = bool(re.search(r"[.!?]$", line))

        if not current:
            current.append(line)
        elif is_degree or has_pipe:
            current.append(line)
        elif is_very_short and not has_sent_end:
            current[0] = current[0] + ", " + line
        elif _looks_like_institution(line):
            blocks.append("\n".join(current))
            current = [line]
        else:
            current.append(line)

    if current:
        blocks.append("\n".join(current))

    return blocks if blocks else ["\n".join(lines)]
